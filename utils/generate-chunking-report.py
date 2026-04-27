#!/usr/bin/env python3
"""Generate a Word doc visualizing logical chunks alongside their source recursive chunks.

Reads pairs of files from a chunking-output directory:
  - {doc_id}-chunks.json        (recursive intermediate chunks)
  - {doc_id}-logic-chunks.json  (logical chunks with `source_chunk_ids` provenance)

Emits one .docx per pair where each logical chunk is shown in a 2-column table
next to the recursive chunks it was built from.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt, RGBColor


LOGIC_SUFFIX = "-logic-chunks.json"
RECURSIVE_SUFFIX = "-chunks.json"
BODY_FONT_PT = 9.5
LABEL_FONT_PT = 9.0
COL_WIDTH = Cm(9.0)
OVERLAP_COLOR = RGBColor(0xA0, 0xA0, 0xA0)  # grey for recursive-chunk overlap regions


def discover_doc_ids(input_dir: Path) -> list[str]:
    return sorted(p.name[: -len(LOGIC_SUFFIX)] for p in input_dir.glob(f"*{LOGIC_SUFFIX}"))


def load_pair(input_dir: Path, doc_id: str) -> tuple[dict, dict] | None:
    rec_path = input_dir / f"{doc_id}{RECURSIVE_SUFFIX}"
    log_path = input_dir / f"{doc_id}{LOGIC_SUFFIX}"
    if not rec_path.exists():
        print(f"[skip] {doc_id}: missing {rec_path.name}", file=sys.stderr)
        return None
    if not log_path.exists():
        print(f"[skip] {doc_id}: missing {log_path.name}", file=sys.stderr)
        return None
    return json.loads(rec_path.read_text()), json.loads(log_path.read_text())


def write_label(cell, text: str) -> None:
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(LABEL_FONT_PT)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def write_chunk_text(cell, text: str) -> None:
    """Render chunk text — no special coloring."""
    write_styled_chunk_text(cell, [(text, None)])


def write_styled_chunk_text(cell, segments: list[tuple[str, "RGBColor | None"]]) -> None:
    """Render a chunk made up of colored segments.

    `segments` is a list of (text, color) tuples in document order; color=None means default.
    `\\n\\n` splits paragraphs (so Word can break cleanly across pages); `\\n` becomes a
    line break within a paragraph.
    """
    paragraphs: list[list[tuple[str, "RGBColor | None"]]] = [[]]
    for seg_text, seg_color in segments:
        parts = seg_text.split("\n\n")
        for k, part in enumerate(parts):
            if k > 0:
                paragraphs.append([])
            if part:
                paragraphs[-1].append((part, seg_color))

    for para_runs in paragraphs:
        para = cell.add_paragraph()
        for txt, color in para_runs:
            lines = txt.split("\n")
            for j, line in enumerate(lines):
                if j > 0:
                    para.add_run().add_break()
                if not line:
                    continue
                run = para.add_run(line)
                run.font.size = Pt(BODY_FONT_PT)
                if color is not None:
                    run.font.color.rgb = color


def longest_overlap(a: str, b: str) -> int:
    """Length of the longest suffix of `a` that equals a prefix of `b`."""
    max_k = min(len(a), len(b))
    for k in range(max_k, 0, -1):
        if a.endswith(b[:k]):
            return k
    return 0


def split_for_overlap(
    text: str, head_len: int, tail_len: int
) -> list[tuple[str, "RGBColor | None"]]:
    """Split `text` into [head_overlap | middle | tail_overlap] colored segments."""
    if head_len + tail_len > len(text):
        # overlap regions collide — clamp the tail so middle doesn't go negative
        tail_len = max(0, len(text) - head_len)
    head = text[:head_len]
    tail = text[len(text) - tail_len:] if tail_len else ""
    middle = text[head_len: len(text) - tail_len] if tail_len else text[head_len:]
    segments: list[tuple[str, "RGBColor | None"]] = []
    if head:
        segments.append((head, OVERLAP_COLOR))
    if middle:
        segments.append((middle, None))
    if tail:
        segments.append((tail, OVERLAP_COLOR))
    return segments


def build_report(rec_data: dict, log_data: dict, out_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    doc_id = log_data.get("doc_id") or rec_data.get("doc_id") or out_path.stem
    rec_by_id = {c["chunk_id"]: c for c in rec_data["texts"]}
    log_chunks = log_data["texts"]

    rec_total_tokens = sum(c.get("tokens", 0) for c in rec_data["texts"])
    log_total_tokens = sum(c.get("tokens", 0) for c in log_chunks)

    doc.add_heading(f"Chunking Report — {doc_id}", level=0)

    summary = doc.add_paragraph()
    parsed_label = summary.add_run("parsed_file: ")
    parsed_label.bold = True
    summary.add_run(str(log_data.get("parsed_file", "")))
    summary.add_run().add_break()
    summary.add_run(f"recursive chunks: {len(rec_data['texts'])} (total tokens: {rec_total_tokens})")
    summary.add_run().add_break()
    summary.add_run(f"logical chunks:   {len(log_chunks)} (total tokens: {log_total_tokens})")
    summary.add_run().add_break()
    legend = summary.add_run(
        "Note: recursive chunks include `chunk_overlap` at boundaries; "
        "the duplicated head/tail regions in the right column are shown in grey "
        "and are removed when merged into a logical chunk."
    )
    legend.italic = True
    legend.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    for idx, lc in enumerate(log_chunks):
        source_ids = lc.get("source_chunk_ids", [])
        heading = (
            f"Logical Chunk {lc['chunk_id']}  ·  tokens={lc.get('tokens', '?')}  "
            f"·  sources={source_ids}"
        )
        doc.add_heading(heading, level=2)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.style = "Table Grid"
        for col in table.columns:
            col.width = COL_WIDTH
        for cell in table.rows[0].cells:
            cell.width = COL_WIDTH
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        left, right = table.rows[0].cells

        write_label(left, f"Logical chunk #{lc['chunk_id']} (tokens={lc.get('tokens', '?')})")
        write_chunk_text(left, lc.get("text", ""))

        # precompute pairwise overlaps between consecutive sources of THIS logical chunk
        source_texts: list[str | None] = [
            (rec_by_id[sid]["text"] if sid in rec_by_id else None) for sid in source_ids
        ]
        overlaps: list[int] = []  # overlaps[i] = overlap between source i and source i+1
        for i in range(len(source_texts) - 1):
            a, b = source_texts[i], source_texts[i + 1]
            overlaps.append(longest_overlap(a, b) if a is not None and b is not None else 0)

        for i, sid in enumerate(source_ids):
            rc = rec_by_id.get(sid)
            if i > 0:
                right.add_paragraph()
            if rc is None:
                write_label(right, f"Recursive #{sid} (MISSING)")
                continue
            head_len = overlaps[i - 1] if i > 0 else 0
            tail_len = overlaps[i] if i < len(overlaps) else 0
            label = f"Recursive #{sid} (tokens={rc.get('tokens', '?')}"
            if head_len or tail_len:
                parts = []
                if head_len:
                    parts.append(f"head overlap {head_len} chars")
                if tail_len:
                    parts.append(f"tail overlap {tail_len} chars")
                label += f"; {', '.join(parts)}"
            label += ")"
            write_label(right, label)
            segments = split_for_overlap(rc.get("text", ""), head_len, tail_len)
            write_styled_chunk_text(right, segments)

        if idx < len(log_chunks) - 1:
            doc.add_page_break()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument(
        "--input-dir", required=True, type=Path,
        help="Directory containing both -chunks.json and -logic-chunks.json files",
    )
    ap.add_argument(
        "--output-dir", type=Path, default=None,
        help="Where to write .docx files (default: <input-dir>/reports)",
    )
    ap.add_argument(
        "--doc-id", default=None,
        help="Process only this doc_id (default: all pairs in input-dir)",
    )
    args = ap.parse_args()

    input_dir: Path = args.input_dir
    if not input_dir.is_dir():
        print(f"input-dir does not exist: {input_dir}", file=sys.stderr)
        return 2
    output_dir: Path = args.output_dir or (input_dir / "reports")

    doc_ids = [args.doc_id] if args.doc_id else discover_doc_ids(input_dir)
    if not doc_ids:
        print(f"no *{LOGIC_SUFFIX} files found in {input_dir}", file=sys.stderr)
        return 1

    n_ok = 0
    for doc_id in doc_ids:
        pair = load_pair(input_dir, doc_id)
        if pair is None:
            continue
        rec, log = pair
        out_path = output_dir / f"{doc_id}-chunking-report.docx"
        build_report(rec, log, out_path)
        print(f"[ok] {out_path}")
        n_ok += 1

    print(f"\nwrote {n_ok}/{len(doc_ids)} report(s) to {output_dir}")
    return 0 if n_ok == len(doc_ids) else 1


if __name__ == "__main__":
    sys.exit(main())
